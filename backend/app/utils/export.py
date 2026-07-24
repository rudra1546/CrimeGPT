import io
import json
import logging
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

logger = logging.getLogger("crimegpt.utils.export")

def load_system_config():
    import os
    data_dir = os.path.join(
        os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
        "data"
    )
    filepath = os.path.join(data_dir, "system_config.json")
    if not os.path.exists(filepath):
        return {
            "government_name": "Government of India",
            "department_name": "State Police Department",
            "logo_url": "/logo_placeholder.png",
            "address": "Administrative Block, State Police Headquarters, India"
        }
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {
            "government_name": "Government of India",
            "department_name": "State Police Department",
            "logo_url": "/logo_placeholder.png",
            "address": "Administrative Block, State Police Headquarters, India"
        }

import re

def strip_md(text: str) -> str:
    if not text or not isinstance(text, str):
        return text or ""
    s = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    s = re.sub(r"\*(.*?)\*", r"\1", s)
    s = re.sub(r"__(.*?)__", r"\1", s)
    s = re.sub(r"_(.*?)_", r"\1", s)
    s = s.replace("**", "").replace("__", "")
    s = re.sub(r"^\s*#{1,6}\s*", "", s)
    s = re.sub(r"^\s*[\*\-]\s*([A-Za-z0-9_\s\/]+:)", r"\1", s)
    return s.strip()

def parse_document_content(content: str) -> dict:
    """
    Parses document content. If it's a valid JSON representation, returns it.
    Otherwise, wraps plain text in a standard structured document format.
    """
    cleaned = content.strip()
    if cleaned.startswith("```"):
        first_line_end = cleaned.find("\n")
        if first_line_end != -1:
            cleaned = cleaned[first_line_end:].strip()
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3].strip()

    try:
        data = json.loads(cleaned)
        if isinstance(data, dict):
            if data.get("document_title"):
                data["document_title"] = strip_md(data["document_title"])
            if isinstance(data.get("meta_info"), dict):
                data["meta_info"] = {strip_md(k): strip_md(str(v)) for k, v in data["meta_info"].items()}
            if isinstance(data.get("body_sections"), list):
                for sec in data["body_sections"]:
                    if isinstance(sec, dict):
                        if sec.get("heading"):
                            sec["heading"] = strip_md(sec["heading"])
                        text_k = "content" if "content" in sec else "text"
                        if sec.get(text_k):
                            sec[text_k] = strip_md(sec[text_k])
            if isinstance(data.get("signatories"), list):
                for sig in data["signatories"]:
                    if isinstance(sig, dict):
                        if sig.get("label"):
                            sig["label"] = strip_md(sig["label"])
                        if sig.get("value"):
                            sig["value"] = strip_md(sig["value"])
            return data
    except Exception:
        pass

    # Fallback format for plain text or legacy records
    return {
        "document_title": "CASE INVESTIGATION RECORD / REPORT",
        "meta_info": {},
        "body_sections": [
            {
                "heading": "DETAILED CASE REPORT",
                "text": strip_md(content)
            }
        ],
        "signatories": [
            {
                "label": "AUTHORIZED SIGNATURE",
                "value": "Investigating Officer / Command Unit"
            }
        ]
    }

def generate_pdf_bytes(content: str) -> io.BytesIO:
    """
    Converts plain text or structured JSON case content into a premium,
    formal government-style PDF report using ReportLab.
    """
    doc_data = parse_document_content(content)
    buffer = io.BytesIO()
    
    # Predefined document margins and template details
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Custom premium styles matching clean government theme
    title_style = ParagraphStyle(
        'GovTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#0f172a'), # Navy blue
        alignment=1, # Centered
        spaceAfter=15
    )
    
    section_heading = ParagraphStyle(
        'GovSectionHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=15,
        textColor=colors.HexColor('#1e3a8a'), # Indigo navy
        spaceBefore=10,
        spaceAfter=5
    )
    
    body_style = ParagraphStyle(
        'GovBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=14,
        textColor=colors.HexColor('#334155'), # Slate text
        spaceAfter=8
    )
    
    meta_style = ParagraphStyle(
        'GovMetaText',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#1e293b')
    )
    
    meta_bold_style = ParagraphStyle(
        'GovMetaTextBold',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#1e293b')
    )

    story = []
    
    # 1. Government-Style Letterhead Header
    config = load_system_config()
    gov_text = f"{config.get('government_name', 'Government of India')} — {config.get('department_name', 'State Police Department')}".upper()
    story.append(Paragraph(gov_text, title_style))
    if config.get("address"):
        addr_style = ParagraphStyle(
            'GovAddrText',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=7,
            alignment=1,
            textColor=colors.HexColor('#6b7280')
        )
        story.append(Paragraph(config.get("address").upper(), addr_style))
    
    # Sleek line separator
    line_table = Table([[""]], colWidths=[504], rowHeights=[2])
    line_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#1e3a8a')),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
        ('TOPPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(line_table)
    story.append(Spacer(1, 10))
    
    # 2. Document Title
    doc_title = doc_data.get("document_title", "OFFICIAL CASE RECORD").upper()
    story.append(Paragraph(doc_title, title_style))
    story.append(Spacer(1, 10))
    
    # 3. Meta-Information Table (Predefined Fixed Grid Layout)
    meta_info = doc_data.get("meta_info", {})
    if meta_info:
        meta_rows = []
        meta_keys = list(meta_info.keys())
        # Render 2 columns of key-value pairs
        for i in range(0, len(meta_keys), 2):
            k1 = meta_keys[i]
            v1 = meta_info[k1]
            p1_key = Paragraph(f"<b>{k1}:</b>", meta_bold_style)
            p1_val = Paragraph(str(v1), meta_style)
            
            p2_key, p2_val = "", ""
            if i + 1 < len(meta_keys):
                k2 = meta_keys[i+1]
                v2 = meta_info[k2]
                p2_key = Paragraph(f"<b>{k2}:</b>", meta_bold_style)
                p2_val = Paragraph(str(v2), meta_style)
                
            meta_rows.append([p1_key, p1_val, p2_key, p2_val])
            
        if meta_rows:
            meta_table = Table(meta_rows, colWidths=[100, 150, 100, 154])
            meta_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
                ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                ('TOPPADDING', (0,0), (-1,-1), 4),
            ]))
            story.append(meta_table)
            story.append(Spacer(1, 15))
            
    # 4. Structured Document Body Sections
    for sec in doc_data.get("body_sections", []):
        heading = sec.get("heading", "")
        text = sec.get("text", sec.get("content", ""))
        if heading:
            story.append(Paragraph(heading.upper(), section_heading))
        # Support paragraphs inside text block
        for para in text.split("\n\n"):
            if para.strip():
                formatted_para = para.strip().replace("\n", "<br/>")
                story.append(Paragraph(formatted_para, body_style))
                
    story.append(Spacer(1, 20))
    
    # 5. Signatories Grid Placement (Keep Together to prevent split across pages)
    signatories = doc_data.get("signatories", [])
    if signatories:
        sig_cols = []
        for sig in signatories:
            label = sig.get("label", "Signature")
            val = sig.get("value", "")
            sig_html = f"<br/><br/>_______________________<br/><b>{label}</b><br/>{val}"
            sig_cols.append(Paragraph(sig_html, meta_style))
            
        # Distribute columns evenly
        col_width = 504 / max(1, len(sig_cols))
        sig_table = Table([sig_cols], colWidths=[col_width] * len(sig_cols))
        sig_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ]))
        story.append(KeepTogether(sig_table))

    pdf.build(story)
    buffer.seek(0)
    return buffer

def generate_docx_bytes(content: str) -> io.BytesIO:
    """
    Converts plain text or structured JSON content into a formal government-style
    Microsoft Word (.docx) document using python-docx.
    """
    doc_data = parse_document_content(content)
    docx = DocxDocument()
    
    # 1. Header Letterhead Title
    config = load_system_config()
    header_p = docx.add_paragraph()
    gov_text = f"{config.get('government_name', 'Government of India')} — {config.get('department_name', 'State Police Department')}".upper()
    r = header_p.add_run(gov_text)
    r.bold = True
    r.font.size = 14 * 12700  # Approx 14pt (using Word units)
    header_p.alignment = 1  # Center

    if config.get("address"):
        addr_p = docx.add_paragraph()
        r_addr = addr_p.add_run(config.get("address").upper())
        r_addr.font.size = 9 * 12700
        addr_p.alignment = 1
    
    # 2. Document Title
    doc_title = doc_data.get("document_title", "OFFICIAL CASE RECORD").upper()
    title_p = docx.add_paragraph()
    r2 = title_p.add_run(doc_title)
    r2.bold = True
    r2.font.size = 13 * 12700
    title_p.alignment = 1
    
    # Divider line
    docx.add_paragraph("__________________________________________________________________").alignment = 1
    
    # 3. Meta-Information Grid (Word Table)
    meta_info = doc_data.get("meta_info", {})
    if meta_info:
        table = docx.add_table(rows=0, cols=2)
        table.autofit = True
        for k, v in meta_info.items():
            row_cells = table.add_row().cells
            row_cells[0].paragraphs[0].add_run(str(k)).bold = True
            row_cells[1].paragraphs[0].add_run(str(v))
        docx.add_paragraph() # Spacer
        
    # 4. Structured Body Sections
    for sec in doc_data.get("body_sections", []):
        heading = sec.get("heading", "")
        text = sec.get("text", sec.get("content", ""))
        if heading:
            p_heading = docx.add_paragraph()
            r_head = p_heading.add_run(heading.upper())
            r_head.bold = True
            
        for para in text.split("\n\n"):
            if para.strip():
                docx.add_paragraph(para.strip())
                
    docx.add_paragraph()
    
    # 5. Signatories List
    signatories = doc_data.get("signatories", [])
    if signatories:
        sig_table = docx.add_table(rows=1, cols=len(signatories))
        cells = sig_table.rows[0].cells
        for idx, sig in enumerate(signatories):
            label = sig.get("label", "Signature")
            val = sig.get("value", "")
            p_sig = cells[idx].paragraphs[0]
            p_sig.add_run("\n\n\n_______________________\n").bold = True
            p_sig.add_run(f"{label}\n").bold = True
            p_sig.add_run(str(val))
            
    buffer = io.BytesIO()
    docx.save(buffer)
    buffer.seek(0)
    return buffer
