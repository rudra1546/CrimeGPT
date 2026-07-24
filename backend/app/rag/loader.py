import os
import json
import logging
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = logging.getLogger("crimegpt.rag.loader")

def load_pdf_text(file_path: str) -> str:
    import pypdf
    reader = pypdf.PdfReader(file_path)
    text_content = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_content.append(page_text)
    return "\n".join(text_content)

def split_text_to_chunks(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> list[Document]:
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        add_start_index=True
    )
    return text_splitter.create_documents([text])

def load_pdf_documents(file_path: str, filename: str) -> list[Document]:
    import pypdf
    reader = pypdf.PdfReader(file_path)
    documents = []
    total_text_length = 0
    for idx, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text and page_text.strip():
            total_text_length += len(page_text)
            meta = {
                "source": filename,
                "file_type": "pdf",
                "page": idx + 1
            }
            documents.append(Document(page_content=page_text.strip(), metadata=meta))
    
    logger.info(f"[Loader] Extracted {total_text_length} chars from {len(reader.pages)} PDF pages for '{filename}'.")
    if not documents:
        return []

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents(documents)
    logger.info(f"[Loader] Created {len(chunks)} chunks for PDF '{filename}'.")
    return chunks

def load_csv_documents(file_path: str, filename: str) -> list[Document]:
    import csv
    documents = []
    with open(file_path, "r", encoding="utf-8-sig", errors="ignore") as f:
        reader = csv.DictReader(f)
        for idx, row in enumerate(reader):
            lines = []
            act_val = "Bharatiya Nyaya Sanhita, 2023" if "bns" in filename.lower() else None
            year_val = "2023" if "bns" in filename.lower() else None
            section_val = None
            offence_val = None
            chapter_val = None
            punishment_val = None
            description_val = None

            for key, val in row.items():
                if not key or not val or not str(val).strip():
                    continue
                k_clean = str(key).strip()
                v_clean = str(val).strip()
                lines.append(f"{k_clean}:\n{v_clean}\n")
                
                kl = k_clean.lower()
                if kl in ["section", "sec", "section_no", "section_number", "code"]:
                    section_val = str(v_clean).strip()
                elif kl in ["section_name", "offence", "offense", "crime", "title", "crime_type"]:
                    offence_val = v_clean
                elif kl in ["act", "act_name", "statute", "law"]:
                    act_val = v_clean
                elif kl in ["year", "act_year"]:
                    year_val = v_clean
                elif kl in ["chapter", "chapter_no"]:
                    chapter_val = v_clean
                elif kl in ["punishment", "penalty", "provision"]:
                    punishment_val = v_clean
                elif kl in ["description", "details", "summary"]:
                    description_val = v_clean

            if lines:
                text_content = "\n".join(lines)
                meta = {
                    "source": filename,
                    "file_type": "csv",
                    "row": idx + 1
                }
                if section_val: meta["section"] = str(section_val).strip()
                if act_val: meta["act_name"] = act_val
                if year_val: meta["year"] = year_val
                if offence_val: meta["offence"] = offence_val
                if chapter_val: meta["chapter"] = chapter_val
                if punishment_val: meta["punishment"] = punishment_val
                if description_val: meta["description"] = description_val

                documents.append(Document(page_content=text_content, metadata=meta))

    logger.info(f"[Loader] Extracted {len(documents)} structured CSV row documents for '{filename}'.")
    return documents

def load_docx_documents(file_path: str, filename: str) -> list[Document]:
    import docx
    doc = docx.Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                paragraphs.append(row_text)
                
    full_text = "\n\n".join(paragraphs)
    logger.info(f"[Loader] Extracted {len(full_text)} chars from DOCX '{filename}'.")
    if not full_text:
        return []
        
    doc_obj = Document(
        page_content=full_text,
        metadata={"source": filename, "file_type": "docx"}
    )
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents([doc_obj])
    logger.info(f"[Loader] Created {len(chunks)} chunks for DOCX '{filename}'.")
    return chunks

def load_txt_documents(file_path: str, filename: str) -> list[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read().strip()
    logger.info(f"[Loader] Extracted {len(text)} chars from TXT '{filename}'.")
    if not text:
        return []
    doc_obj = Document(
        page_content=text,
        metadata={"source": filename, "file_type": "txt"}
    )
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents([doc_obj])
    logger.info(f"[Loader] Created {len(chunks)} chunks for TXT '{filename}'.")
    return chunks

def load_md_documents(file_path: str, filename: str) -> list[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read().strip()
    logger.info(f"[Loader] Extracted {len(text)} chars from Markdown '{filename}'.")
    if not text:
        return []
    doc_obj = Document(
        page_content=text,
        metadata={"source": filename, "file_type": "md"}
    )
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents([doc_obj])
    logger.info(f"[Loader] Created {len(chunks)} chunks for Markdown '{filename}'.")
    return chunks

def load_json_documents(file_path: str, filename: str) -> list[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        data = json.load(f)
    
    documents = []
    if isinstance(data, list):
        for idx, item in enumerate(data):
            if isinstance(item, dict):
                lines = []
                act_val = "Bharatiya Nyaya Sanhita, 2023" if "bns" in filename.lower() else None
                year_val = "2023" if "bns" in filename.lower() else None
                section_val = None
                offence_val = None
                for k, v in item.items():
                    if v and str(v).strip():
                        k_clean = str(k).strip()
                        v_clean = str(v).strip()
                        lines.append(f"{k_clean}:\n{v_clean}\n")
                        kl = k_clean.lower()
                        if kl in ["section", "sec", "section_no", "code"]: section_val = str(v_clean).strip()
                        elif kl in ["section_name", "offence", "offense", "title"]: offence_val = v_clean
                        elif kl in ["act", "act_name", "law"]: act_val = v_clean
                        elif kl in ["year"]: year_val = v_clean

                if lines:
                    meta = {"source": filename, "file_type": "json", "item": idx + 1}
                    if section_val: meta["section"] = str(section_val).strip()
                    if act_val: meta["act_name"] = act_val
                    if year_val: meta["year"] = year_val
                    if offence_val: meta["offence"] = offence_val
                    documents.append(Document(page_content="\n".join(lines), metadata=meta))
    elif isinstance(data, dict):
        lines = []
        for k, v in data.items():
            lines.append(f"{k}:\n{json.dumps(v, indent=2)}\n")
        documents.append(Document(page_content="\n".join(lines), metadata={"source": filename, "file_type": "json"}))
        
    logger.info(f"[Loader] Extracted {len(documents)} JSON element documents for '{filename}'.")
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents(documents) if documents else []
    logger.info(f"[Loader] Created {len(chunks)} chunks for JSON '{filename}'.")
    return chunks

def load_and_chunk_file(file_path: str, filename: str) -> list[Document]:
    """
    Auto-detects format (PDF, CSV, DOCX, TXT, JSON, Markdown) and extracts structured documents with metadata.
    """
    ext = os.path.splitext(filename)[1].lower()
    logger.info(f"[Loader] Processing uploaded document '{filename}' (type: {ext})...")
    if ext == ".pdf":
        return load_pdf_documents(file_path, filename)
    elif ext == ".csv":
        return load_csv_documents(file_path, filename)
    elif ext in [".docx", ".doc"]:
        return load_docx_documents(file_path, filename)
    elif ext == ".txt":
        return load_txt_documents(file_path, filename)
    elif ext in [".md", ".markdown"]:
        return load_md_documents(file_path, filename)
    elif ext == ".json":
        return load_json_documents(file_path, filename)
    else:
        raise ValueError(f"Unsupported format '{ext}'. Supported formats: PDF, CSV, DOCX, TXT, JSON, MD.")
