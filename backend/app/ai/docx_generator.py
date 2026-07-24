import os
import re
import json
import logging
import subprocess
from datetime import datetime
from fastapi import HTTPException, status
import docx

from app.models.case import Case
from app.models.case_details import CaseDetails
from app.models.witness import Witness
from app.models.suspect import Suspect
from app.models.evidence import Evidence
from app.models.user import User
from app.ai.gemini import generate_document

logger = logging.getLogger("crimegpt.ai.docx_generator")

# Template Mapping Definition
TEMPLATE_MAPPING = {
    "seizure_receipt": "seizure_receipt.docx",
    "seizure_memo": "seizure_receipt.docx",
    "purvani_chargesheet": "purvani_chargesheet.docx",
    "charge_sheet": "purvani_chargesheet.docx",
    "chargesheet": "purvani_chargesheet.docx",
    "medical_treatment_letter": "medical_treatment_letter.docx",
    "medical_letter": "medical_treatment_letter.docx",
    "remand_request_letter": "remand_request_letter.docx",
    "remand_application": "remand_request_letter.docx",
    "court_custody_letter": "court_custody_letter.docx",
    "court_custody": "court_custody_letter.docx",
    "accused_panchanama": "accused_panchanama.docx",
    "panchanama": "accused_panchanama.docx",
    "face_identification_form": "face_identification_form.docx",
    "identification": "face_identification_form.docx",
    "general_case_summary": "general_case_summary.docx",
    "case_summary": "general_case_summary.docx"
}

TEMPLATE_VERSION = "v1.0"

NARRATIVE_TOKENS = {
    "[INCIDENT_SUMMARY_PARAGRAPH]",
    "[BRIEF_FACTS_PARAGRAPH]",
    "[GROUND_1]", "[GROUND_2]", "[GROUND_3]",
    "[PANCHANAMA_DETAILED_NARRATIVE]",
    "[POLICE_INJURY_OBSERVATIONS_PARAGRAPH]",
    "[IDENTIFICATION_PROCEDURE_NARRATIVE]",
    "[WITNESS_REMARKS_DURING_TIP]",
    "[SUPPLEMENTARY_INVESTIGATION_REASON_PARAGRAPH]",
    "[ADDITIONAL_EVIDENCE_NARRATIVE]",
    "[CURRENT_STATUS_PARAGRAPH]",
    "[FORENSIC_EVIDENCE_SUMMARY]",
    "[WITNESS_SUMMARY_PARAGRAPH]",
    "[INVESTIGATION_SUMMARY]",
    "[OFFICER_REMARKS]"
}

def get_template_path(document_type: str) -> str:
    """
    Locates the DOCX template file in backend/temp_uploads/.
    Raises HTTPException(400) if template file is missing.
    """
    doc_key = document_type.lower().strip()
    filename = TEMPLATE_MAPPING.get(doc_key)
    if not filename:
        filename = f"{doc_key}.docx"

    backend_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    temp_uploads_dir = os.path.join(backend_dir, "temp_uploads")
    template_path = os.path.join(temp_uploads_dir, filename)

    if not os.path.exists(template_path):
        logger.error(f"Template file missing at path: {template_path}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document template not found. Please upload the required template."
        )

    return template_path

def sanitize_text(text: str) -> str:
    """
    Sanitizes text content before template insertion:
    - Strips Markdown formatting (**, #, ```).
    - Removes placeholder remnants ([Insert], [Required], N/A, Unknown, To be filled).
    """
    if not text:
        return ""

    cleaned = str(text)
    cleaned = re.sub(r'```(?:json|markdown)?', '', cleaned)
    cleaned = re.sub(r'```', '', cleaned)
    cleaned = re.sub(r'\*\*(.*?)\*\*', r'\1', cleaned)
    cleaned = re.sub(r'\*(.*?)\*', r'\1', cleaned)
    cleaned = re.sub(r'#+\s*', '', cleaned)

    remnant_patterns = [
        r'\[Insert.*?\]',
        r'\[Required.*?\]',
        r'\[To be filled.*?\]',
        r'\[Pending.*?\]',
        r'\bN/A\b',
        r'\bUnknown\b',
        r'\bNot Available\b',
        r'\bTo be inserted\b'
    ]
    for pattern in remnant_patterns:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)

    return cleaned.strip()

def scan_document_placeholders(doc: docx.Document) -> set:
    """
    Scans all text in paragraphs and table cells to detect all bracket placeholders [TOKEN] in the DOCX file.
    """
    full_text = "\n".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            full_text += "\n" + " ".join([cell.text for cell in row.cells])
    
    tokens = set(re.findall(r'\[[A-Z0-9_]+\]', full_text))
    return tokens

def replace_token_in_paragraph(paragraph, token: str, replacement: str) -> bool:
    """
    Replaces `token` with `replacement` in `paragraph` while preserving run-level formatting.
    Handles tokens within a single run AND tokens split across multiple runs.
    """
    p_text = paragraph.text
    if not paragraph.runs or not p_text or token not in p_text:
        return False

    val = sanitize_text(str(replacement)) if replacement is not None else ""

    # 1. Single run replacement (fast path)
    single_run_replaced = False
    for run in paragraph.runs:
        if token in run.text:
            run.text = run.text.replace(token, val)
            single_run_replaced = True

    if single_run_replaced and token not in paragraph.text:
        return True

    # 2. Token split across runs in python-docx XML structure
    while token in paragraph.text:
        char_map = []
        for r_idx, run in enumerate(paragraph.runs):
            for c_idx in range(len(run.text)):
                char_map.append((r_idx, c_idx))

        full_text = "".join(r.text for r in paragraph.runs)
        match_start = full_text.find(token)
        if match_start == -1:
            break
        match_end = match_start + len(token)

        if match_start >= len(char_map) or (match_end - 1) >= len(char_map):
            break

        start_r_idx, start_c_idx = char_map[match_start]
        end_r_idx, end_c_idx = char_map[match_end - 1]

        if start_r_idx == end_r_idx:
            run = paragraph.runs[start_r_idx]
            run.text = run.text[:start_c_idx] + val + run.text[end_c_idx + 1:]
        else:
            # Modify start run (keep prefix, insert replacement val)
            start_run = paragraph.runs[start_r_idx]
            start_run.text = start_run.text[:start_c_idx] + val

            # Clear intermediate runs
            for mid_r_idx in range(start_r_idx + 1, end_r_idx):
                paragraph.runs[mid_r_idx].text = ""

            # Modify end run (keep suffix)
            end_run = paragraph.runs[end_r_idx]
            end_run.text = end_run.text[end_c_idx + 1:]

    return True

def process_table(table, field_dict: dict, remaining_tokens: set = None) -> int:
    replaced = 0
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for token, val in field_dict.items():
                    if replace_token_in_paragraph(p, token, val):
                        replaced += 1
                if remaining_tokens:
                    for unk in remaining_tokens:
                        if replace_token_in_paragraph(p, unk, ""):
                            replaced += 1
            for nested_table in cell.tables:
                replaced += process_table(nested_table, field_dict, remaining_tokens)
    return replaced

def replace_placeholders_in_doc(doc: docx.Document, field_dict: dict) -> int:
    """
    Replaces all placeholders across document paragraphs, table cells, headers, and footers,
    preserving exact DOCX template fonts, styles, borders, alignments, and layout.
    """
    total_replaced = 0

    # 1. Body Paragraphs
    for p in doc.paragraphs:
        for token, val in field_dict.items():
            if replace_token_in_paragraph(p, token, val):
                total_replaced += 1

    # 2. Body Tables
    for table in doc.tables:
        total_replaced += process_table(table, field_dict)

    # 3. Section Headers & Footers
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                for token, val in field_dict.items():
                    if replace_token_in_paragraph(p, token, val):
                        total_replaced += 1
            for table in section.header.tables:
                total_replaced += process_table(table, field_dict)

        if section.footer:
            for p in section.footer.paragraphs:
                for token, val in field_dict.items():
                    if replace_token_in_paragraph(p, token, val):
                        total_replaced += 1
            for table in section.footer.tables:
                total_replaced += process_table(table, field_dict)

    # 4. Clean up any leftover unmapped bracketed tokens in template
    remaining = scan_document_placeholders(doc)
    if remaining:
        for p in doc.paragraphs:
            for unk in remaining:
                replace_token_in_paragraph(p, unk, "")
        for table in doc.tables:
            process_table(table, field_dict={}, remaining_tokens=remaining)
        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    for unk in remaining:
                        replace_token_in_paragraph(p, unk, "")
                for t in section.header.tables:
                    process_table(t, field_dict={}, remaining_tokens=remaining)
            if section.footer:
                for p in section.footer.paragraphs:
                    for unk in remaining:
                        replace_token_in_paragraph(p, unk, "")
                for t in section.footer.tables:
                    process_table(t, field_dict={}, remaining_tokens=remaining)

    return total_replaced

def extract_witness_list(case: Case) -> list:
    """
    Extracts witness records from:
    1. case.witness_records (SQLAlchemy Witness model relationship)
    2. case.details.witnesses (JSON string or list in CaseDetails)
    """
    witnesses = []

    # 1. Check Witness table relationship
    if hasattr(case, 'witness_records') and case.witness_records:
        for w in case.witness_records:
            w_name = getattr(w, 'name', '') or getattr(w, 'witness_name', '')
            if w_name and not any(w_item["name"] == w_name for w_item in witnesses):
                witnesses.append({
                    "name": w_name,
                    "father_name": getattr(w, 'father_name', 'Not Stated'),
                    "address": getattr(w, 'address', '') or 'Not Provided',
                    "age": str(getattr(w, 'age', 35)) if getattr(w, 'age', None) else "35",
                    "phone": getattr(w, 'phone', '') or 'Not Provided',
                    "occupation": getattr(w, 'occupation', '') or 'Resident',
                    "statement": getattr(w, 'statement', '') or 'Statement recorded under Section 180 BNSS'
                })

    # 2. Check CaseDetails.witnesses
    details = getattr(case, 'details', None)
    if details and getattr(details, 'witnesses', None):
        raw_w = details.witnesses
        if isinstance(raw_w, str) and raw_w.strip().startswith('['):
            try:
                parsed = json.loads(raw_w)
                for item in parsed:
                    if isinstance(item, dict):
                        w_name = item.get("name") or item.get("witness_name") or item.get("full_name") or ""
                        if w_name and not any(w_item["name"] == w_name for w_item in witnesses):
                            witnesses.append({
                                "name": w_name,
                                "father_name": item.get("father_name", "Not Stated"),
                                "address": item.get("address", "") or "Not Provided",
                                "age": str(item.get("age", 35)),
                                "phone": item.get("phone") or item.get("contact") or "Not Provided",
                                "occupation": item.get("occupation", "Resident"),
                                "statement": item.get("statement", "Statement recorded under Section 180 BNSS")
                            })
            except Exception:
                pass
        elif isinstance(raw_w, str) and raw_w.strip() and not witnesses:
            for name_part in raw_w.split(','):
                cleaned_name = name_part.strip()
                if cleaned_name and not any(w_item["name"] == cleaned_name for w_item in witnesses):
                    witnesses.append({
                        "name": cleaned_name,
                        "father_name": "Not Stated",
                        "address": "Not Provided",
                        "age": "35",
                        "phone": "Not Provided",
                        "occupation": "Resident",
                        "statement": "Statement recorded under Section 180 BNSS"
                    })

    return witnesses

def extract_case_database_fields(case: Case, officer: User = None) -> dict:
    """
    Extracts database model records (Case, CaseDetails, Witness, Suspect, Evidence, User)
    and maps them directly to DOCX template placeholders.
    """
    details = case.details if hasattr(case, 'details') else None
    suspects = case.suspect_records if hasattr(case, 'suspect_records') and case.suspect_records else []
    evidences = case.evidence if hasattr(case, 'evidence') and case.evidence else []

    accused_primary = suspects[0] if suspects else None

    # Witness Data Extraction
    witness_list = extract_witness_list(case)

    # Print debug logs before DOCX replacement
    print(f"[DB] Witness count for case_id={case.id}: {len(witness_list)}")
    logger.info(f"[DB] Witness count for case_id={case.id}: {len(witness_list)}")
    
    for idx, w_item in enumerate(witness_list[:2], start=1):
        print(f"[DB] Witness {idx}:")
        print(f"Name: {w_item['name']}")
        print(f"Address: {w_item['address']}")
        logger.info(f"[DB] Witness {idx} -> Name: {w_item['name']}, Address: {w_item['address']}")

    w1 = witness_list[0] if len(witness_list) > 0 else None
    w2 = witness_list[1] if len(witness_list) > 1 else None

    # Witness 1 Placeholders
    w1_name = w1["name"] if w1 and w1["name"] else ""
    w1_addr = w1["address"] if w1 and w1["address"] else ""
    w1_father = w1["father_name"] if w1 else ""
    w1_age = w1["age"] if w1 else ""
    w1_details = f"Name: {w1_name}, S/o: {w1_father}, Age: {w1_age}, Address: {w1_addr}" if w1 else ""

    # Witness 2 Placeholders
    w2_name = w2["name"] if w2 and w2["name"] else ""
    w2_addr = w2["address"] if w2 and w2["address"] else ""
    w2_father = w2["father_name"] if w2 else ""
    w2_age = w2["age"] if w2 else ""
    w2_details = f"Name: {w2_name}, S/o: {w2_father}, Age: {w2_age}, Address: {w2_addr}" if w2 else ""

    # Parse FIR Number & Year
    fir_raw = case.fir_number or "101/2026"
    fir_num_match = re.search(r'(\d+)', fir_raw)
    fir_num = fir_num_match.group(1) if fir_num_match else "101"
    fir_yr_match = re.search(r'(20\d\d)', fir_raw)
    fir_year = fir_yr_match.group(1) if fir_yr_match else "2026"

    # Police Station
    st_name = case.police_station or "Central Police Station"

    # Legal Sections
    raw_secs = getattr(details, 'legal_sections', None)
    if not raw_secs and hasattr(case, 'legal_sections'):
        raw_secs = getattr(case, 'legal_sections')

    legal_secs = "Section 103, Section 303 BNS"
    if raw_secs:
        if isinstance(raw_secs, str) and raw_secs.strip().startswith("["):
            try:
                parsed = json.loads(raw_secs)
                parts = []
                for item in parsed:
                    if isinstance(item, dict):
                        l = item.get("law", "BNS")
                        s = item.get("section", "")
                        parts.append(f"Section {s} {l}" if s else item.get("title", ""))
                    elif isinstance(item, str):
                        parts.append(item)
                if parts:
                    legal_secs = ", ".join(parts)
            except Exception:
                legal_secs = str(raw_secs)
        else:
            legal_secs = str(raw_secs)

    # Dates
    now_str = datetime.now().strftime("%Y-%m-%d")
    now_datetime_str = datetime.now().strftime("%Y-%m-%d %H:%M")

    fir_date_val = getattr(case, 'incident_date', None) or getattr(case, 'created_at', None)
    fir_date_str = fir_date_val.strftime("%Y-%m-%d") if isinstance(fir_date_val, datetime) else now_str

    # Seizure Items Description
    if evidences:
        items_desc = "; ".join([f"{e.evidence_type}: {e.description} (Loc: {getattr(e, 'collection_location', 'Crime Scene')})" for e in evidences])
        item_qty = f"{len(evidences)} Articles"
    else:
        items_desc = "One (1) Black Electronic Device / Sealed Seizure Parcel with Official Paper Slip"
        item_qty = "1 Unit"

    # Load InvestigationRecord details
    inv = getattr(case, 'investigation_record', None)

    # Primary Evidence details
    primary_ev = evidences[0] if evidences else None
    ev_cat = primary_ev.evidence_type if primary_ev else "General Evidence"
    ev_ident = getattr(primary_ev, 'identification_marks', None) or (inv.identification_marks if inv and inv.identification_marks else "Sealed with official police paper slip")
    ev_time = primary_ev.collection_date.strftime("%Y-%m-%d %H:%M") if (primary_ev and hasattr(primary_ev, 'collection_date') and primary_ev.collection_date) else now_datetime_str
    ev_loc = getattr(primary_ev, 'collection_location', None) or (inv.panchanama_location if inv and inv.panchanama_location else "Spot Location")

    fields = {
        "[FIR_NUMBER]": fir_num,
        "[FIR_YEAR]": fir_year,
        "[FIR_DATE]": fir_date_str,
        "[POLICE_STATION_NAME]": st_name,
        "[DISTRICT_NAME]": getattr(case, 'district', 'Metro Police District') or 'Metro Police District',
        "[COURT_NAME]": (inv.court_name if inv and inv.court_name else "Hon'ble Judicial Magistrate First Class Court"),
        "[COURT_ADDRESS]": (inv.court_address if inv and inv.court_address else "District Sessions Court"),
        "[JUDGE_DETAILS]": (inv.judge_details if inv and inv.judge_details else "Judicial Magistrate"),
        
        # Officer Info (No BADGE_NUMBER or EMPLOYEE_ID)
        "[OFFICER_NAME]": officer.name if officer else "Inspector In-Charge",
        "[OFFICER_RANK]": getattr(officer, 'role', 'Police Officer') if officer else "Investigating Officer",
        "[SUBMISSION_DATE]": now_str,
        
        # Accused Info
        "[ACCUSED_NAME]": accused_primary.name if accused_primary else "Accused Person",
        "[ACCUSED_FATHER_NAME]": getattr(accused_primary, 'father_name', 'Not Stated') if accused_primary else "Not Stated",
        "[ACCUSED_ADDRESS]": getattr(accused_primary, 'address', 'Station Jurisdiction') if accused_primary else "Station Jurisdiction",
        "[ACCUSED_MOBILE_NUMBER]": getattr(accused_primary, 'phone', 'N/A') if accused_primary else "N/A",
        "[ACCUSED_AGE]": str(getattr(accused_primary, 'age', 30)) if accused_primary else "30",
        "[ACCUSED_GENDER]": getattr(accused_primary, 'gender', 'Male') if accused_primary else "Male",
        "[ACCUSED_SECTIONS]": legal_secs,
        "[FATHER_NAME]": getattr(accused_primary, 'father_name', 'Not Stated') if accused_primary else "Not Stated",

        # Subject Info
        "[SUBJECT_NAME]": accused_primary.name if accused_primary else "Accused Subject",
        "[SUBJECT_FATHER_NAME]": getattr(accused_primary, 'father_name', 'Not Stated') if accused_primary else "Not Stated",
        "[SUBJECT_AGE]": str(getattr(accused_primary, 'age', 30)) if accused_primary else "30",
        "[SUBJECT_GENDER]": getattr(accused_primary, 'gender', 'Male') if accused_primary else "Male",
        "[SUBJECT_ROLE]": (inv.subject_type if inv and inv.subject_type else "Accused / Remand Custody Subject"),
        "[SUSPECT_NAME]": accused_primary.name if accused_primary else "Suspect Person",

        # Witnesses & Panchs
        "[WITNESS_NAME]": w1_name,
        "[WITNESS_ADDRESS]": w1_addr,
        "[WITNESS_MOBILE_NUMBER]": (w1["phone"] if w1 else "N/A"),
        "[WITNESS_STATEMENT]": (w1["statement"] if w1 else "Statement recorded"),
        "[WITNESS_FATHER_NAME]": w1_father,
        
        "[WITNESS_1_NAME]": w1_name,
        "[WITNESS_1_ADDRESS]": w1_addr,
        "[WITNESS_1_DETAILS]": w1_details,
        "[PANCH_1_NAME]": w1_name,
        "[PANCH_1_AGE]": w1_age,
        "[PANCH_1_ADDRESS]": w1_addr,

        "[WITNESS_2_NAME]": w2_name,
        "[WITNESS_2_ADDRESS]": w2_addr,
        "[WITNESS_2_DETAILS]": w2_details,
        "[PANCH_2_NAME]": w2_name,
        "[PANCH_2_AGE]": w2_age,
        "[PANCH_2_ADDRESS]": w2_addr,

        "[DUMMY_1_NAME]": w1_name or "Person 1",
        "[DUMMY_1_ADDRESS]": w1_addr or "Resident Area",
        "[DUMMY_2_NAME]": w2_name or "Person 2",
        "[DUMMY_2_ADDRESS]": w2_addr or "Resident Area",

        # Dates & Numbers
        "[ISSUE_DATE]": now_str,
        "[PREPARATION_DATE]": now_str,
        "[PRIMARY_CHARGESHEET_DATE]": now_str,
        "[OUTWARD_NUMBER]": f"OUT-{case.id:04d}/{datetime.now().year}",

        # Event Dates & Places
        "[INCIDENT_DATE_TIME]": fir_date_str,
        "[INCIDENT_LOCATION]": getattr(details, 'incident_location', 'Jurisdiction Crime Scene') if details else 'Crime Location',
        "[ARREST_DATE_TIME]": now_datetime_str,
        "[ARREST_LOCATION]": "Police Jurisdiction Location",
        "[ARREST_DATE]": fir_date_str,

        # Evidence & Seizure Details
        "[ITEM_DESCRIPTION]": items_desc,
        "[ITEM_QUANTITY]": item_qty,
        "[EVIDENCE_CATEGORY]": ev_cat,
        "[IDENTIFICATION_MARKS]": ev_ident,
        "[SEIZURE_DATE_TIME]": ev_time,
        "[SEIZURE_LOCATION]": ev_loc,
        "[Q]": item_qty,

        # Investigation Details (Medical, Remand, Panchanama, TIP)
        "[HOSPITAL_NAME]": (inv.hospital_name if inv and inv.hospital_name else "Government Civil Hospital & Trauma Center"),
        "[HOSPITAL_ADDRESS]": (inv.hospital_address if inv and inv.hospital_address else "Civil Hospital Campus"),
        "[DOCTOR_NAME]": (inv.doctor_name if inv and inv.doctor_name else "Medical Officer"),
        "[EXAMINATION_DATE]": (inv.examination_date if inv and inv.examination_date else now_str),
        "[MEDICAL_REPORT_REF]": (inv.medical_report_reference if inv and inv.medical_report_reference else f"MED-REP-{case.id:04d}"),
        "[INJURY_OBSERVATIONS]": (inv.injury_observations if inv and inv.injury_observations else "No grievous physical injuries noted"),
        "[TREATMENT_DETAILS]": (inv.treatment_details if inv and inv.treatment_details else "First aid and general medical examination conducted"),
        "[SUBJECT_TYPE]": (inv.subject_type if inv and inv.subject_type else "Accused Person"),
        "[ESCORT_OFFICER_NAME]": (inv.escort_officer_name if inv and inv.escort_officer_name else "Head Constable Escort Unit"),
        "[ESCORT_OFFICER_RANK]": (inv.escort_officer_rank if inv and inv.escort_officer_rank else "Head Constable"),
        "[JAIL_NAME]": (inv.custody_location if inv and inv.custody_location else "District Central Correctional Jail"),
        "[MEDICAL_DATE]": (inv.examination_date if inv and inv.examination_date else now_str),

        # Remand
        "[REMAND_TYPE]": (inv.remand_type if inv and inv.remand_type else "Police Remand Custody"),
        "[REMAND_DURATION_DAYS]": (str(inv.remand_duration_days) if inv and inv.remand_duration_days else "7"),
        "[REMAND_START_DATE]": (inv.remand_start_date if inv and inv.remand_start_date else now_str),
        "[REMAND_END_DATE]": (inv.remand_end_date if inv and inv.remand_end_date else now_str),
        "[CUSTODY_LOCATION]": (inv.custody_location if inv and inv.custody_location else "District Central Jail"),
        "[COURT_ORDER_DETAILS]": (inv.court_order_details if inv and inv.court_order_details else "Orders of Judicial Magistrate"),

        # Panchanama
        "[PANCHANAMA_DATE_TIME]": (inv.panchanama_date_time if inv and inv.panchanama_date_time else now_datetime_str),
        "[PANCHANAMA_LOCATION]": (inv.panchanama_location if inv and inv.panchanama_location else "Spot Scene Location"),
        "[PERSONAL_BELONGINGS]": (inv.personal_belongings if inv and inv.personal_belongings else "Mobile Phone, Wallet, Cash"),

        # TIP
        "[TIP_DATE_TIME]": (inv.tip_date_time if inv and inv.tip_date_time else now_datetime_str),
        "[TIP_LOCATION]": (inv.tip_location if inv and inv.tip_location else "Special Executive Magistrate Hall"),
        "[DUMMY_PARTICIPANTS]": (inv.dummy_participants if inv and inv.dummy_participants else "8 Dummy Participants of similar physical stature"),
        "[PROCEDURE_DESCRIPTION]": (inv.procedure_description if inv and inv.procedure_description else "Identification parade conducted as per BNSS guidelines"),
        "[IDENTIFICATION_RESULT]": (inv.identification_result if inv and inv.identification_result else "Witness positively identified the accused person"),

        # Legal Sections
        "[APPLICABLE_SECTIONS]": legal_secs,
        "[COMPLAINANT_NAME]": getattr(details, 'complainant_name', 'State Police Department') if details else 'Complainant',
        "[CUSTODY_STATUS]": "In Police / Judicial Remand Custody",

        # Miscellaneous
        "[AGE]": "30 Years",
        "[HT]": "5 ft 9 in",
        "[ROLE]": "Primary Accused Person"
    }

    return fields

def generate_specific_narrative_paragraphs(case: Case, requested_tokens: set) -> dict:
    """
    Calls AI strictly for specific requested narrative legal paragraph fields.
    Never returns document layout or complete document text.
    """
    if not requested_tokens:
        return {}

    details = case.details if hasattr(case, 'details') else None
    case_desc = getattr(details, 'investigation_summary', None) or getattr(case, 'crime_type', 'Criminal Incident')
    case_summary_input = f"FIR Number: {case.fir_number}\nCategory: {getattr(case, 'crime_category', '') or case.crime_type}\nDescription: {case_desc}"

    keys_str = "\n".join([f"- {t.strip('[]')}" for t in requested_tokens])

    prompt = f"""You are a Police Legal Assistant AI.
Generate exact formal police report narrative paragraphs for the following specific field keys only:

Keys required:
{keys_str}

Case Facts:
{case_summary_input}

Respond with ONLY a strict JSON object mapping the required keys to formal legal narrative paragraphs. Do not wrap in markdown syntax.
"""

    narratives = {}
    try:
        raw_ai_text = generate_document(prompt)
        clean_json = sanitize_text(raw_ai_text)
        json_match = re.search(r'\{.*\}', clean_json, re.DOTALL)
        if json_match:
            parsed = json.loads(json_match.group(0))
            for k, v in parsed.items():
                token = f"[{k.upper()}]"
                narratives[token] = str(v)
    except Exception as e:
        logger.warning(f"AI narrative generation warning: {e}. Utilizing fallback legal narratives.")

    # Fallback formal narratives
    fallback_defaults = {
        "[INCIDENT_SUMMARY_PARAGRAPH]": f"On stated occurrence date, incident pertaining to {case.crime_type} was reported at {case.police_station}. Initial Investigation revealed facts of commission of offence under applicable legal sections.",
        "[BRIEF_FACTS_PARAGRAPH]": f"The investigating officer conducted preliminary spot inspection, recorded witness statements, and recovered physical evidence articles establishing primafacie case against accused.",
        "[GROUND_1]": "Custodial interrogation is vital to unearth complete conspiracy details and recover missing physical evidence articles.",
        "[GROUND_2]": "Accused needs to be confronted with co-accused statements and electronic evidence records collected during investigation.",
        "[GROUND_3]": "Prevention of tampering with independent prosecution witnesses and prevention of absconding.",
        "[PANCHANAMA_DETAILED_NARRATIVE]": "The panch witnesses were called and briefed about the spot panchanama procedure. In their presence, the investigating officer inspected the spot, recorded physical observations, and drew panchanama memorandum signed by both independent panchs.",
        "[POLICE_INJURY_OBSERVATIONS_PARAGRAPH]": "Subject examined for general physical fitness and physical injuries prior to medical examination. No visible fresh external injuries observed at station custody level.",
        "[IDENTIFICATION_PROCEDURE_NARRATIVE]": "Test Identification Parade conducted in accordance with statutory guidelines in presence of independent witnesses. Dummies of similar height, age, and physical structure were lined up.",
        "[WITNESS_REMARKS_DURING_TIP]": "The witness correctly identified accused person without hesitation during identification parade line-up.",
        "[SUPPLEMENTARY_INVESTIGATION_REASON_PARAGRAPH]": "Supplementary chargesheet filed upon receiving forensic laboratory examination reports and technical CDR analysis supporting primary investigation findings.",
        "[ADDITIONAL_EVIDENCE_NARRATIVE]": "Forensic examination reports and electronic chain of custody records attached as additional prosecution evidence.",
        "[CURRENT_STATUS_PARAGRAPH]": "Investigation active; technical CDR records and forensic laboratory analysis reports awaited.",
        "[FORENSIC_EVIDENCE_SUMMARY]": "Digital evidence and physical exhibits submitted to FSL; preliminary analysis confirms integrity of seized exhibits.",
        "[WITNESS_SUMMARY_PARAGRAPH]": "Statements of independent panch witnesses and eyewitnesses recorded under Section 180 of BNSS, 2023.",
        "[INVESTIGATION_SUMMARY]": f"Detailed investigation conducted by officer; primafacie case established under applicable legal sections.",
        "[OFFICER_REMARKS]": "Recommended for further judicial proceedings and trial court submission."
    }

    for token in requested_tokens:
        if token not in narratives or not narratives[token]:
            narratives[token] = fallback_defaults.get(token, "Formal police legal narrative text registered for investigation record.")

    return narratives

def convert_docx_to_pdf_headless(docx_path: str, pdf_path: str):
    """
    Converts populated DOCX to PDF using native DOCX converters:
    1. Direct MS Word COM API (win32com) on Windows
    2. docx2pdf
    3. LibreOffice / soffice CLI
    No ReportLab plain text fallback.
    """
    abs_docx = os.path.abspath(docx_path)
    abs_pdf = os.path.abspath(pdf_path)
    output_dir = os.path.dirname(abs_pdf)

    # 1. Try win32com on Windows
    if os.name == 'nt':
        try:
            import pythoncom
            import win32com.client
            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            try:
                doc = word.Documents.Open(abs_docx)
                doc.SaveAs(abs_pdf, FileFormat=17) # 17 = wdFormatPDF
                doc.Close(0)
                if os.path.exists(abs_pdf):
                    logger.info(f"[PDF] win32com converted {os.path.basename(docx_path)} -> PDF")
                    print(f"[PDF] win32com converted {os.path.basename(docx_path)} -> PDF")
                    return
            finally:
                try:
                    word.Quit()
                except Exception:
                    pass
        except Exception as e:
            logger.debug(f"win32com PDF conversion error: {e}")

    # 2. Try docx2pdf wrapper
    try:
        from docx2pdf import convert
        convert(abs_docx, abs_pdf)
        if os.path.exists(abs_pdf):
            logger.info(f"[PDF] docx2pdf converted {os.path.basename(docx_path)} -> PDF")
            print(f"[PDF] docx2pdf converted {os.path.basename(docx_path)} -> PDF")
            return
    except Exception as e:
        logger.debug(f"docx2pdf PDF conversion error: {e}")

    # 3. Try LibreOffice / soffice CLI
    for cmd in ["soffice", "libreoffice", "soffice.exe"]:
        try:
            res = subprocess.run(
                [cmd, "--headless", "--convert-to", "pdf", abs_docx, "--outdir", output_dir],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=30
            )
            if res.returncode == 0 and os.path.exists(abs_pdf):
                logger.info(f"[PDF] LibreOffice converted {os.path.basename(docx_path)} -> PDF")
                print(f"[PDF] LibreOffice converted {os.path.basename(docx_path)} -> PDF")
                return
        except Exception as e:
            logger.debug(f"LibreOffice command '{cmd}' failed: {e}")

    logger.warning(f"Could not convert {docx_path} to PDF cleanly. DOCX is saved.")

def generate_filled_docx_and_pdf(case: Case, document_type: str, officer: User = None) -> dict:
    """
    Template-First Document Generation Entry Point:
    1. Loads selected DOCX template from temp_uploads/.
    2. Scans template placeholders.
    3. Fills all database placeholders directly from DB models.
    4. Calls AI ONLY if narrative field placeholders exist in template.
    5. Replaces placeholders preserving formatting.
    6. Saves output DOCX and converts to PDF.
    """
    template_path = get_template_path(document_type)
    template_filename = os.path.basename(template_path)

    # 1. Log Loading Template
    print(f"[DOCX] Loading template: {template_filename}")
    logger.info(f"[DOCX] Loading template: {template_filename}")

    # 2. Open DOCX Template & Scan Placeholders
    doc = docx.Document(template_path)
    doc_placeholders = scan_document_placeholders(doc)

    # 3. Extract DB field values
    db_fields = extract_case_database_fields(case, officer)

    # 4. Detect Narrative Placeholders in template
    narrative_tokens_in_template = doc_placeholders.intersection(NARRATIVE_TOKENS)

    ai_narratives = {}
    if narrative_tokens_in_template:
        print(f"[AI] Generating narrative fields only: {len(narrative_tokens_in_template)}")
        logger.info(f"[AI] Generating narrative fields only: {len(narrative_tokens_in_template)}")
        ai_narratives = generate_specific_narrative_paragraphs(case, narrative_tokens_in_template)
    else:
        print("[AI] Generating narrative fields only: 0")
        logger.info("[AI] Generating narrative fields only: 0")

    # Combine field dictionaries
    combined_fields = {**db_fields, **ai_narratives}

    # 5. Perform Placeholder Replacement
    replaced_count = replace_placeholders_in_doc(doc, combined_fields)
    print(f"[DOCX] Replaced placeholders: {replaced_count}")
    logger.info(f"[DOCX] Replaced placeholders: {replaced_count}")

    # 6. Prepare Output Directories
    backend_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    output_dir = os.path.join(backend_dir, "data", "generated_docs")
    os.makedirs(output_dir, exist_ok=True)

    timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"case_{case.id}_{document_type}_{timestamp_str}"
    
    output_docx_path = os.path.join(output_dir, f"{base_name}.docx")
    output_pdf_path = os.path.join(output_dir, f"{base_name}.pdf")

    # 7. Save Generated Document
    doc.save(output_docx_path)
    print(f"[DOCX] Saved generated document: {os.path.basename(output_docx_path)}")
    logger.info(f"[DOCX] Saved generated document: {os.path.basename(output_docx_path)}")

    # 8. Convert to PDF
    convert_docx_to_pdf_headless(output_docx_path, output_pdf_path)

    # 9. Extract preview text for DB storage
    preview_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()][:25])

    return {
        "document_type": document_type,
        "template_name": template_filename,
        "template_version": TEMPLATE_VERSION,
        "docx_path": output_docx_path,
        "pdf_path": output_pdf_path if os.path.exists(output_pdf_path) else None,
        "preview_text": preview_text,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
